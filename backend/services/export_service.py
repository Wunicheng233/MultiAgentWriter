"""
导出服务 - 支持多格式导出小说
- EPUB: 电子书格式
- DOCX: Word 格式
- HTML: 静态网页格式
"""

import os
import datetime
import re
from html import escape, unescape
from html.parser import HTMLParser
from typing import Dict, Tuple
from pathlib import Path

from sqlalchemy.orm import Session
from backend.models import Project, Chapter


SAFE_EXPORT_TAGS = {
    'a',
    'blockquote',
    'br',
    'code',
    'div',
    'em',
    'h1',
    'h2',
    'h3',
    'h4',
    'h5',
    'h6',
    'hr',
    'li',
    'ol',
    'p',
    'pre',
    's',
    'strong',
    'table',
    'tbody',
    'td',
    'th',
    'thead',
    'tr',
    'u',
    'ul',
}
VOID_EXPORT_TAGS = {'br', 'hr'}
BLOCKED_EXPORT_TAGS = {'script', 'style', 'iframe', 'object', 'embed'}
SAFE_LINK_PROTOCOLS = {'http', 'https', 'mailto'}


class SafeExportHTMLParser(HTMLParser):
    """Small allowlist sanitizer for exported HTML documents."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.blocked_depth = 0

    def handle_starttag(self, tag: str, attrs):
        tag = tag.lower()
        if tag in BLOCKED_EXPORT_TAGS:
            self.blocked_depth += 1
            return
        if self.blocked_depth or tag not in SAFE_EXPORT_TAGS:
            return

        safe_attrs = self._safe_attrs(tag, attrs)
        attrs_text = ''.join(f' {name}="{escape(value, quote=True)}"' for name, value in safe_attrs)
        self.parts.append(f'<{tag}{attrs_text}>')

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in BLOCKED_EXPORT_TAGS:
            self.blocked_depth = max(0, self.blocked_depth - 1)
            return
        if self.blocked_depth or tag not in SAFE_EXPORT_TAGS or tag in VOID_EXPORT_TAGS:
            return
        self.parts.append(f'</{tag}>')

    def handle_startendtag(self, tag: str, attrs):
        tag = tag.lower()
        if self.blocked_depth or tag not in VOID_EXPORT_TAGS:
            return
        self.parts.append(f'<{tag}>')

    def handle_data(self, data: str):
        if not self.blocked_depth:
            self.parts.append(escape(data))

    def _safe_attrs(self, tag: str, attrs) -> list[tuple[str, str]]:
        safe_attrs: list[tuple[str, str]] = []
        attrs_dict = {name.lower(): value for name, value in attrs if value is not None}

        if tag == 'a':
            href = attrs_dict.get('href', '').strip()
            if self._is_safe_href(href):
                safe_attrs.append(('href', href))
                safe_attrs.append(('target', '_blank'))
                safe_attrs.append(('rel', 'noopener noreferrer'))
            title = attrs_dict.get('title', '').strip()
            if title:
                safe_attrs.append(('title', title))

        if tag in {'td', 'th'}:
            for attr_name in ('colspan', 'rowspan'):
                attr_value = attrs_dict.get(attr_name, '').strip()
                if attr_value.isdigit():
                    safe_attrs.append((attr_name, attr_value))

        return safe_attrs

    @staticmethod
    def _is_safe_href(href: str) -> bool:
        if not href:
            return False
        if href.startswith(('/', '#')):
            return True
        protocol_match = re.match(r'^([a-zA-Z][a-zA-Z0-9+.-]*):', href)
        return bool(protocol_match and protocol_match.group(1).lower() in SAFE_LINK_PROTOCOLS)

    def get_html(self) -> str:
        return ''.join(self.parts)


class ExportService:
    """多格式导出服务"""

    def __init__(self, db: Session, project_id: int):
        self.db = db
        self.project_id = project_id
        from sqlalchemy.orm import joinedload
        self.project = db.query(Project)\
            .options(joinedload(Project.owner))\
            .filter(Project.id == project_id)\
            .first()
        self.chapters = db.query(Chapter)\
            .filter(Chapter.project_id == project_id)\
            .order_by(Chapter.chapter_index)\
            .all()

    def get_project_info(self) -> Dict:
        """获取项目基本信息"""
        config = self.project.config or {}
        return {
            'title': self._clean_text(config.get('novel_name', self.project.name)),
            'description': self._clean_text(config.get('novel_description', self.project.description or '')),
            'author': self._clean_text(self.project.owner.username if self.project.owner else 'Anonymous'),
            'chapters': [
                {
                    'index': c.chapter_index,
                    'title': self._clean_text(c.title or f'第{c.chapter_index}章'),
                    'content': self._clean_html_content(c.content),
                    'word_count': c.word_count,
                } for c in self.chapters
            ]
        }

    def _clean_text(self, value: str) -> str:
        """Return display text with any embedded HTML removed."""
        text = str(value or '')
        if self._looks_like_html(text):
            return self._strip_html(self._clean_html_content(text))
        return text

    def _clean_html_content(self, content: str) -> str:
        """清理 TipTap HTML 内容，转换为纯文本/干净HTML"""
        if not content:
            return ''
        if not self._looks_like_html(content):
            return ''.join(f'<p>{escape(line.strip())}</p>' for line in content.splitlines() if line.strip())

        parser = SafeExportHTMLParser()
        parser.feed(content)
        parser.close()
        return parser.get_html()

    def _safe_filename_component(self, value: str) -> str:
        """Return a path-safe filename component without leaking path separators."""
        normalized = re.sub(r'[^\w.\-\u4e00-\u9fff]+', '_', value, flags=re.UNICODE).strip('._-')
        return normalized[:80] or f'project_{self.project_id}'

    @staticmethod
    def _looks_like_html(content: str) -> bool:
        return bool(re.search(r'</?[a-zA-Z][^>]*>', content))

    def export_epub(self, output_dir: str = '/tmp') -> Tuple[str, str]:
        """
        导出 EPUB 格式
        返回 (file_path, filename)
        """
        try:
            from ebooklib import epub
        except ImportError:
            raise RuntimeError("ebooklib not installed, please install with: pip install ebooklib")

        info = self.get_project_info()
        book = epub.EpubBook()

        # 设置元数据
        book.set_identifier(f'storyforge-{self.project_id}-{int(datetime.datetime.utcnow().timestamp())}')
        book.set_title(info['title'])
        book.set_language('zh-CN')
        book.add_author(info['author'])

        # 添加章节
        epub_chapters = []
        for chap in info['chapters']:
            c = epub.EpubHtml(
                title=chap['title'],
                file_name=f'chap_{chap["index"]:02d}.xhtml',
                lang='zh-CN'
            )
            c.content = f'<h1>{escape(chap["title"])}</h1>\n{chap["content"]}'
            book.add_item(c)
            epub_chapters.append(c)

        # 添加导航
        book.toc = epub_chapters
        book.spine = ['nav'] + epub_chapters

        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # 生成文件
        os.makedirs(output_dir, exist_ok=True)
        filename = f'{self.project_id}_{self._safe_filename_component(info["title"])}_{datetime.datetime.now():%Y%m%d}.epub'
        output_path = os.path.join(output_dir, filename)
        epub.write_epub(output_path, book, {})

        return output_path, filename

    def export_docx(self, output_dir: str = '/tmp') -> Tuple[str, str]:
        """
        导出 DOCX 格式
        返回 (file_path, filename)
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise RuntimeError("python-docx not installed, please install with: pip install python-docx")

        info = self.get_project_info()
        doc = Document()

        # 标题
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(info['title'])
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        # 作者
        if info['author']:
            author_paragraph = doc.add_paragraph()
            author_run = author_paragraph.add_run(f'作者：{info["author"]}')
            author_run.font.size = Pt(14)
            author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()

        # 描述
        if info['description']:
            doc.add_paragraph(info['description'])
            doc.add_paragraph()

        # 添加章节
        for chap in info['chapters']:
            # 章节标题
            heading = doc.add_heading(chap['title'], level=1)
            # 移除默认格式影响，设置居中
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 章节内容，需要去掉 HTML 标签转换为纯文本
            # 简单处理：去掉 HTML 标签
            content_text = self._strip_html(chap['content'])
            # 分段
            paragraphs = content_text.split('\n\n')
            for p in paragraphs:
                if p.strip():
                    doc.add_paragraph(p.strip())

            doc.add_page_break()

        # 保存文件
        os.makedirs(output_dir, exist_ok=True)
        filename = f'{self.project_id}_{self._safe_filename_component(info["title"])}_{datetime.datetime.now():%Y%m%d}.docx'
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

        return output_path, filename

    def export_html(self, output_dir: str = '/tmp') -> Tuple[str, str]:
        """
        导出 HTML 单页格式
        返回 (file_path, filename)
        """
        info = self.get_project_info()

        html_template = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escape(info['title'])}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: "Crimson Pro", Georgia, serif;
            font-size: 18px;
            line-height: 1.8;
            color: #3a2c1f;
            background: #faf7f2;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        h1 {{
            font-size: 2.5em;
            text-align: center;
            margin-bottom: 0.5em;
        }}
        .info {{
            text-align: center;
            color: #7a6f62;
            margin-bottom: 3em;
            padding-bottom: 2em;
            border-bottom: 1px solid #e8ddd0;
        }}
        .chapter {{
            margin-bottom: 3em;
        }}
        .chapter h2 {{
            font-size: 1.8em;
            margin-bottom: 1em;
            color: #3a2c1f;
        }}
        p {{
            margin-bottom: 1em;
            text-indent: 2em;
        }}
    </style>
</head>
<body>
    <h1>{escape(info['title'])}</h1>
    <div class="info">
        <p>作者：{escape(info['author'])}</p>
        {f'<p>{escape(info["description"])}</p>' if info['description'] else ''}
    </div>
'''

        for chap in info['chapters']:
            html_template += f'''<div class="chapter">
    <h2>{escape(chap['title'])}</h2>
    {chap['content']}
</div>
'''

        html_template += '''</body></html>'''

        os.makedirs(output_dir, exist_ok=True)
        filename = f'{self.project_id}_{self._safe_filename_component(info["title"])}_{datetime.datetime.now():%Y%m%d}.html'
        output_path = os.path.join(output_dir, filename)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)

        return output_path, filename

    def _strip_html(self, html: str) -> str:
        """简单去除 HTML 标签"""
        text = re.sub(r'<[^>]+>', '', html)
        # 替换多个空行为单个
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return unescape(text).strip()

    @staticmethod
    def cleanup_old_files(directory: str, max_age_hours: int = 24) -> None:
        """清理过期的临时导出文件"""
        now = datetime.datetime.now().timestamp()
        for f in Path(directory).glob('*'):
            if f.is_file():
                age = now - f.stat().st_mtime
                if age > max_age_hours * 3600:
                    try:
                        f.unlink()
                    except OSError:
                        pass
